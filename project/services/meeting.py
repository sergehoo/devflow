"""
DevFlow — Services Meeting (séries, occurrences, IA, compte-rendu).

API exposée par ``MeetingService`` :

  * ``generate_occurrences(series, horizon_days=60)`` : crée les
    ``ProjectMeeting`` futures pour cette série (idempotent grâce à
    ``unique_together(series, scheduled_at)``)
  * ``sync_review_slots(meeting)`` : pour chaque projet dans
    ``meeting.projects`` (ou ``series.default_projects``), crée le
    ``MeetingProjectReview`` correspondant s'il manque
  * ``generate_ai_summary(meeting)`` : synthèse via AIProvider (fallback
    DeepSeek → Ollama) — exploite project_reviews + agenda + notes
  * ``render_minutes_docx(meeting)`` : génère un .docx avec branding
    workspace (logo + tagline + pied de page légal). Réutilise le
    pattern de la facture.
  * ``send_minutes_email(meeting)`` : envoie le .docx + résumé HTML
    aux internal_participants + external_participants

Toutes les opérations sont scopées au workspace de la réunion — aucune
fuite cross-tenant possible.
"""

from __future__ import annotations

import calendar
import io
import logging
from datetime import date, datetime, time, timedelta
from typing import Iterable

from django.conf import settings
from django.contrib.auth import get_user_model
from django.core.mail import EmailMessage
from django.db import transaction
from django.template.loader import render_to_string
from django.utils import timezone

from project import models as dm

logger = logging.getLogger(__name__)
User = get_user_model()


# ---------------------------------------------------------------------------
# Génération d'occurrences
# ---------------------------------------------------------------------------
def _next_weekly_dates(
    start: date, end: date, weekday: int, step_weeks: int = 1,
) -> list[date]:
    """Liste de toutes les dates entre start et end qui tombent un weekday donné."""
    out: list[date] = []
    if weekday is None:
        return out
    # Aligne sur le 1er weekday >= start
    days_ahead = (weekday - start.weekday()) % 7
    cursor = start + timedelta(days=days_ahead)
    while cursor <= end:
        out.append(cursor)
        cursor += timedelta(weeks=step_weeks)
    return out


def _next_daily_workday_dates(start: date, end: date) -> list[date]:
    """Lundi → vendredi entre start et end inclus."""
    out: list[date] = []
    cursor = start
    while cursor <= end:
        if cursor.weekday() < 5:
            out.append(cursor)
        cursor += timedelta(days=1)
    return out


def _next_monthly_dates(start: date, end: date, day_of_month: int) -> list[date]:
    """1er du mois (ou 'dernier' si day_of_month=0) entre start et end."""
    out: list[date] = []
    if day_of_month is None:
        return out
    cursor = date(start.year, start.month, 1)
    while cursor <= end:
        # Détermine le jour réel : si "dernier" (0) ou si jour > max du mois
        max_day = calendar.monthrange(cursor.year, cursor.month)[1]
        if day_of_month == 0:
            real_day = max_day
        else:
            real_day = min(day_of_month, max_day)
        occurrence = date(cursor.year, cursor.month, real_day)
        if start <= occurrence <= end:
            out.append(occurrence)
        # Avance au 1er du mois suivant
        if cursor.month == 12:
            cursor = date(cursor.year + 1, 1, 1)
        else:
            cursor = date(cursor.year, cursor.month + 1, 1)
    return out


class MeetingService:
    """API publique du module réunion."""

    # ─── Génération d'occurrences ─────────────────────────────────────
    @classmethod
    @transaction.atomic
    def generate_occurrences(
        cls,
        series: dm.MeetingSeries,
        *,
        horizon_days: int = 60,
        now: datetime | None = None,
    ) -> list[dm.ProjectMeeting]:
        """
        Crée les ``ProjectMeeting`` à venir pour cette série, jusqu'à
        ``now + horizon_days``. Idempotent.

        Politique :
          * NONE     → une seule occurrence à start_date + time_local
          * DAILY    → tous les jours ouvrés (lun-ven)
          * WEEKLY   → chaque ``weekday`` (1 par semaine)
          * BIWEEKLY → chaque ``weekday`` (1 sur 2 semaines)
          * MONTHLY  → chaque ``month_day`` (0 = dernier jour)

        Si l'occurrence existe déjà (même series + même scheduled_at),
        on ne la duplique pas.
        """
        if not series.is_active:
            return []

        now = now or timezone.localtime()
        today = now.date()
        horizon = today + timedelta(days=horizon_days)
        end = series.end_date or horizon
        end = min(end, horizon)
        start = max(series.start_date, today)

        if start > end:
            return []

        dates: list[date] = []
        if series.recurrence == dm.MeetingSeries.Recurrence.NONE:
            if series.start_date <= horizon:
                dates = [series.start_date]
        elif series.recurrence == dm.MeetingSeries.Recurrence.DAILY:
            dates = _next_daily_workday_dates(start, end)
        elif series.recurrence == dm.MeetingSeries.Recurrence.WEEKLY:
            dates = _next_weekly_dates(start, end, series.weekday or 0, step_weeks=1)
        elif series.recurrence == dm.MeetingSeries.Recurrence.BIWEEKLY:
            dates = _next_weekly_dates(start, end, series.weekday or 0, step_weeks=2)
        elif series.recurrence == dm.MeetingSeries.Recurrence.MONTHLY:
            dates = _next_monthly_dates(start, end, series.month_day or 1)

        created: list[dm.ProjectMeeting] = []
        for d in dates:
            scheduled_dt = timezone.make_aware(
                datetime.combine(d, series.time_local or time(9, 0)),
                timezone.get_current_timezone(),
            )
            # Idempotence : on cherche une occurrence existante
            existing = dm.ProjectMeeting.objects.filter(
                series=series, scheduled_at=scheduled_dt,
            ).first()
            if existing:
                continue
            occ = dm.ProjectMeeting.objects.create(
                workspace=series.workspace,
                series=series,
                title=series.name,
                meeting_type=series.meeting_type,
                status=dm.ProjectMeeting.Status.PLANNED,
                scheduled_at=scheduled_dt,
                duration_minutes=series.duration_minutes,
                location=series.location,
                meeting_link=series.meeting_link,
                organizer=series.organizer,
                agenda=series.default_agenda,
                created_by=series.created_by,
            )
            # Copie les participants par défaut
            participants_ids = list(series.default_participants.values_list("pk", flat=True))
            if participants_ids:
                occ.internal_participants.set(participants_ids)
            # Copie les projets par défaut
            project_ids = list(series.default_projects.values_list("pk", flat=True))
            if project_ids:
                occ.projects.set(project_ids)
                cls.sync_review_slots(occ)
            created.append(occ)
        return created

    # ─── Slots de revue projet ─────────────────────────────────────────
    @classmethod
    @transaction.atomic
    def sync_review_slots(cls, meeting: dm.ProjectMeeting) -> int:
        """
        Pour chaque projet dans ``meeting.projects``, crée le
        ``MeetingProjectReview`` s'il n'existe pas. Retourne le nombre
        de slots ajoutés.

        Sécurité : on ne crée des reviews QUE pour les projets du même
        workspace que la réunion.
        """
        ws = meeting.workspace
        if ws is None:
            return 0
        existing = set(
            meeting.project_reviews.values_list("project_id", flat=True)
        )
        projects = meeting.projects.filter(workspace=ws).exclude(
            pk__in=existing,
        ).only("id")
        max_pos = (
            meeting.project_reviews.order_by("-position").values_list("position", flat=True).first()
            or 0
        )
        new_reviews = []
        for idx, proj in enumerate(projects, start=1):
            new_reviews.append(dm.MeetingProjectReview(
                meeting=meeting, project=proj,
                position=max_pos + idx,
            ))
        if new_reviews:
            dm.MeetingProjectReview.objects.bulk_create(new_reviews)
        return len(new_reviews)

    # ─── Résumé IA (compte-rendu) ──────────────────────────────────────
    @classmethod
    def generate_ai_summary(cls, meeting: dm.ProjectMeeting) -> str:
        """
        Produit une synthèse en Markdown du compte-rendu (agenda + notes
        + revues projet + décisions + actions). Utilise le provider AI
        en mode "auto" (DeepSeek → Ollama fallback silencieux).

        Stocke le résultat dans ``meeting.ai_summary`` et marque
        ``ai_extracted_at``. Retourne le texte produit.
        """
        from project.services.ai.factory import get_ai_provider
        from project.services.ai.base import AIMessage

        provider = get_ai_provider()
        if not provider or not provider.is_available():
            return ""

        reviews = list(meeting.project_reviews.select_related("project"))
        reviews_lines = []
        for r in reviews:
            reviews_lines.append(
                f"### {r.project.name}\n"
                f"- Statut : {r.get_status_snapshot_display()} ({r.progress_pct} %)\n"
                f"- Réalisations : {r.achievements or '—'}\n"
                f"- Bloquants : {r.blockers or '—'}\n"
                f"- Décisions : {r.decisions or '—'}\n"
                f"- Actions à mener : {r.actions_to_take or '—'}\n"
                f"- Prochain jalon : {r.next_milestone or '—'}"
                + (f" ({r.next_milestone_date})" if r.next_milestone_date else "")
            )

        # Liste des participants
        try:
            participants_internal = [
                u.get_full_name() or u.get_username()
                for u in meeting.internal_participants.all()
            ]
        except Exception:
            participants_internal = []

        context_block = (
            f"# Réunion : {meeting.title}\n"
            f"- Type : {meeting.get_meeting_type_display()}\n"
            f"- Date : {timezone.localtime(meeting.scheduled_at).strftime('%d/%m/%Y %H:%M')}\n"
            f"- Durée : {meeting.duration_minutes} min\n"
            f"- Participants internes : "
            f"{', '.join(participants_internal) if participants_internal else '—'}\n"
            f"- Externes : {meeting.external_participants or '—'}\n\n"
            f"## Ordre du jour\n{meeting.agenda or '—'}\n\n"
            f"## Notes brutes\n{meeting.notes or '—'}\n\n"
            f"## Décisions générales\n{meeting.decisions or '—'}\n\n"
            f"## Prochaines étapes globales\n{meeting.next_steps or '—'}\n\n"
            f"## Revue projet par projet\n"
            + ("\n\n".join(reviews_lines) if reviews_lines else "— Aucun projet —")
        )

        system = (
            "Tu es l'assistant DevFlow. Tu rédiges un compte-rendu professionnel, "
            "structuré et concis à partir des notes brutes d'une réunion. "
            "Format Markdown. Sections : Résumé exécutif (3 lignes), Décisions, "
            "Actions (format puces avec @responsable et date si dispo), "
            "Revue par projet (statut + 2 bullets clés par projet), Risques, "
            "Prochaine réunion. Garde les chiffres et noms exacts."
        )

        try:
            response = provider.generate(
                messages=[
                    AIMessage(role="system", content=system),
                    AIMessage(role="user", content=context_block[:8000]),
                ],
                temperature=0.3,
                max_tokens=1500,
            )
            text = (response.text or "").strip()
        except Exception as exc:
            logger.warning("MeetingService.generate_ai_summary failed: %s", exc)
            return ""

        if text:
            meeting.ai_summary = text[:8000]
            meeting.ai_extracted_at = timezone.now()
            meeting.ai_used_provider = response.provider or ""
            meeting.save(update_fields=[
                "ai_summary", "ai_extracted_at", "ai_used_provider", "updated_at",
            ])
        return text

    # ─── Compte-rendu Word (.docx) ─────────────────────────────────────
    @classmethod
    def render_minutes_docx(cls, meeting: dm.ProjectMeeting) -> bytes:
        """
        Génère un compte-rendu Word avec le branding workspace.

        Structure :
          1. En-tête avec logo + tagline workspace
          2. Bloc info réunion (titre, type, date, durée, lieu)
          3. Participants
          4. Ordre du jour
          5. Résumé exécutif IA (si dispo)
          6. Revue projet par projet (1 section par MeetingProjectReview)
          7. Décisions / Actions / Prochaines étapes globales
          8. Pied de page : raison sociale + RCCM + CC + adresse + tél

        Nécessite ``python-docx``. Si non installé, lève ImportError.
        """
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        workspace = meeting.workspace
        accent = (workspace.accent_color or "#F4722B").lstrip("#")
        try:
            accent_rgb = RGBColor(
                int(accent[0:2], 16),
                int(accent[2:4], 16),
                int(accent[4:6], 16),
            )
        except Exception:
            accent_rgb = RGBColor(0xF4, 0x72, 0x2B)

        doc = Document()

        # En-tête (logo + tagline)
        section = doc.sections[0]
        header = section.header
        ht = header.paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if workspace.logo:
            try:
                run = ht.add_run()
                run.add_picture(workspace.logo.path, height=Cm(2.0))
            except Exception:
                pass
        if workspace.tagline:
            ht_tag = header.add_paragraph(workspace.tagline)
            ht_tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in ht_tag.runs:
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = accent_rgb

        # Pied de page (raison sociale + coordonnées)
        footer = section.footer
        footer_lines = []
        if workspace.legal_name or workspace.name:
            footer_lines.append(workspace.legal_name or workspace.name)
        if workspace.legal_rccm:
            footer_lines.append(f"RCCM : {workspace.legal_rccm}")
        if workspace.legal_cc:
            footer_lines.append(f"CC : {workspace.legal_cc}")
        if workspace.address_line1:
            addr = workspace.address_line1
            if workspace.address_line2:
                addr += " — " + workspace.address_line2
            footer_lines.append(addr)
        if workspace.postal_code or workspace.city:
            footer_lines.append(
                f"{workspace.postal_code or ''} {workspace.city or ''}".strip()
            )
        if workspace.phone:
            footer_lines.append(f"Tél. : {workspace.phone}")
        if workspace.email:
            footer_lines.append(workspace.email)
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in footer_lines:
            r = fp.add_run(line + "\n")
            r.font.size = Pt(8)

        # ── Corps ────────────────────────────────────────────────────
        # Titre
        title = doc.add_paragraph()
        run = title.add_run("COMPTE-RENDU DE RÉUNION")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = accent_rgb

        h1 = doc.add_paragraph()
        h1.add_run(meeting.title).bold = True
        h1.runs[0].font.size = Pt(20)

        # Bloc info
        info = doc.add_paragraph()
        info.add_run(f"Type : {meeting.get_meeting_type_display()}\n").bold = True
        info.add_run(
            f"Date : {timezone.localtime(meeting.scheduled_at).strftime('%A %d %B %Y à %H:%M')}\n"
        )
        info.add_run(f"Durée : {meeting.duration_minutes} min\n")
        if meeting.location:
            info.add_run(f"Lieu : {meeting.location}\n")
        if meeting.meeting_link:
            info.add_run(f"Lien : {meeting.meeting_link}\n")

        # Participants
        doc.add_heading("Participants", level=2)
        internal = list(meeting.internal_participants.all())
        if internal:
            p = doc.add_paragraph()
            p.add_run("Internes : ").bold = True
            p.add_run(", ".join(
                u.get_full_name() or u.get_username() for u in internal
            ))
        if meeting.external_participants:
            p = doc.add_paragraph()
            p.add_run("Externes : ").bold = True
            p.add_run(meeting.external_participants)

        # Résumé IA
        if meeting.ai_summary:
            doc.add_heading("Résumé exécutif", level=2)
            for line in meeting.ai_summary.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # Ordre du jour
        if meeting.agenda:
            doc.add_heading("Ordre du jour", level=2)
            for line in meeting.agenda.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")

        # Revue par projet
        reviews = list(meeting.project_reviews.select_related("project")
                       .order_by("position", "id"))
        if reviews:
            doc.add_heading("Revue projet par projet", level=2)
            for r in reviews:
                doc.add_heading(r.project.name, level=3)
                p = doc.add_paragraph()
                p.add_run("Statut : ").bold = True
                p.add_run(f"{r.get_status_snapshot_display()} ({r.progress_pct} %)")
                if r.presented_by:
                    p2 = doc.add_paragraph()
                    p2.add_run("Présenté par : ").bold = True
                    p2.add_run(
                        r.presented_by.get_full_name() or r.presented_by.get_username()
                    )
                for label, content in [
                    ("Réalisations", r.achievements),
                    ("Bloquants", r.blockers),
                    ("Décisions", r.decisions),
                    ("Actions à mener", r.actions_to_take),
                ]:
                    if content:
                        pl = doc.add_paragraph()
                        pl.add_run(label + " : ").bold = True
                        pl.add_run(content)
                if r.next_milestone:
                    pm = doc.add_paragraph()
                    pm.add_run("Prochain jalon : ").bold = True
                    pm.add_run(r.next_milestone)
                    if r.next_milestone_date:
                        pm.add_run(f" ({r.next_milestone_date})")

        # Décisions et next steps globales
        if meeting.decisions:
            doc.add_heading("Décisions générales", level=2)
            doc.add_paragraph(meeting.decisions)
        if meeting.next_steps:
            doc.add_heading("Prochaines étapes", level=2)
            doc.add_paragraph(meeting.next_steps)
        if meeting.blockers:
            doc.add_heading("Bloquants globaux", level=2)
            doc.add_paragraph(meeting.blockers)

        # Notes brutes en annexe
        if meeting.notes:
            doc.add_heading("Notes complémentaires", level=2)
            doc.add_paragraph(meeting.notes)

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    # ─── Envoi email du compte-rendu ───────────────────────────────────
    @classmethod
    def send_minutes_email(cls, meeting: dm.ProjectMeeting,
                           include_external: bool = True) -> int:
        """
        Envoie le .docx + un résumé HTML aux participants. Retourne le
        nombre de destinataires.

        Le sujet inclut le numéro et la date. Le corps est une version
        HTML simple du ``ai_summary`` si dispo, sinon du contenu brut.
        """
        recipients: list[str] = []
        # Participants internes (emails uniquement valides)
        recipients.extend(
            u.email for u in meeting.internal_participants.all() if u.email
        )
        # Participants externes : on parse les lignes, on extrait les
        # adresses email basiques (chaîne contenant @).
        if include_external and meeting.external_participants:
            import re
            for line in meeting.external_participants.splitlines():
                for match in re.findall(r"[\w\.\-+]+@[\w\.\-]+\.\w+", line):
                    if match not in recipients:
                        recipients.append(match)
        if not recipients:
            return 0

        # Génère le .docx
        try:
            docx_bytes = cls.render_minutes_docx(meeting)
        except Exception as exc:
            logger.warning("Cannot render docx for meeting %s: %s", meeting.pk, exc)
            docx_bytes = None

        # Corps HTML
        body_text = meeting.ai_summary or (
            (meeting.notes or "")
            + "\n\n— Compte-rendu DevFlow"
        )
        subject = f"[CR] {meeting.title} — {timezone.localtime(meeting.scheduled_at).strftime('%d/%m/%Y')}"

        from_email = getattr(settings, "DEFAULT_FROM_EMAIL", None) \
            or getattr(settings, "EMAIL_HOST_USER", None) \
            or "noreply@devflow.local"

        sent = 0
        for to in recipients:
            try:
                msg = EmailMessage(
                    subject=subject,
                    body=body_text,
                    from_email=from_email,
                    to=[to],
                )
                if docx_bytes:
                    filename = f"CR-{meeting.title.replace(' ', '_')}.docx"
                    msg.attach(
                        filename, docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                msg.send(fail_silently=True)
                sent += 1
            except Exception as exc:
                logger.warning("Email send failed to %s: %s", to, exc)
        return sent
