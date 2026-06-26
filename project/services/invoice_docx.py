"""
DevFlow — Génération d'une facture au format Microsoft Word (.docx).

Calque sur ``invoice_pdf.py`` : même contexte (invoice, workspace, lines, totaux,
client, project, period, paiements, notes, mentions légales) mais sortie .docx
via *python-docx*.

Layout :
  - En-tête : logo workspace (à gauche) + tagline (à droite)
  - Titre   : "Facture" + N° en orange + badge statut (1 ligne)
  - Bloc client (fond gris clair)
  - Bandeau info 4 colonnes : Date émission / Échéance / Projet ou Référence / Devise
  - Tableau lignes (Désignation / Qté / PU HT / Total HT)
  - Totaux à droite (Sous-total HT, Remise, TVA, Total TTC, Encaissé, Reste dû)
  - Notes éventuelles
  - Mentions légales / coordonnées bancaires (si renseignées)
  - Footer : coordonnées du workspace sur 3 colonnes

Aucune fuite cross-workspace : la facture passée doit déjà avoir été filtrée
par la vue ; ici on lit uniquement ``invoice.workspace``.
"""

from __future__ import annotations

import io
import logging
from decimal import Decimal

from django.utils.formats import number_format

logger = logging.getLogger(__name__)


# ────────────────────────────────────────────────────────────────────
# Constantes de style (calquées sur le PDF compact)
# ────────────────────────────────────────────────────────────────────
ACCENT_DEFAULT = "F4722B"       # orange Kaydan
TEXT_PRIMARY = "1D1A17"          # noir doux
TEXT_MUTED = "6D655D"            # gris moyen
TEXT_FAINT = "A79F97"            # gris clair
BG_SOFT = "F6F3EF"               # fond bloc client / notes
BORDER_SOFT = "E9E1D8"           # bordures fines

# Couleurs badges (matching PDF)
BADGE_COLORS = {
    "PAID":      {"bg": "DCFCE7", "fg": "15803D", "label": "Payée"},
    "OVERDUE":   {"bg": "FEE2E2", "fg": "B91C1C", "label": "En retard"},
    "DRAFT":     {"bg": "F6F3EF", "fg": "6D655D", "label": "Brouillon"},
    "ISSUED":    {"bg": "DBEAFE", "fg": "1D4ED8", "label": "Émise"},
    "SENT":      {"bg": "DBEAFE", "fg": "1D4ED8", "label": "Envoyée"},
    "CANCELLED": {"bg": "FEE2E2", "fg": "B91C1C", "label": "Annulée"},
}


# ────────────────────────────────────────────────────────────────────
# Helpers de mise en forme
# ────────────────────────────────────────────────────────────────────
def _fmt_amount(value, currency: str = "") -> str:
    if value is None:
        value = Decimal("0")
    try:
        # Format à 2 décimales + séparateur de milliers FR
        formatted = number_format(value, decimal_pos=2, use_l10n=False, force_grouping=True)
    except Exception:
        formatted = f"{Decimal(str(value)):,.2f}"
    if currency:
        return f"{formatted} {currency}"
    return formatted


def _fmt_date(d):
    if not d:
        return "—"
    try:
        return d.strftime("%d/%m/%Y")
    except Exception:
        return str(d)


def _hex_to_rgb(hex_color: str):
    from docx.shared import RGBColor
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    """Force une couleur de fond sur une cellule."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None, size="4"):
    """
    Ajoute des bordures sur une cellule.
    Chaque paramètre est soit None (pas de bordure), soit un dict
    {"sz": "8", "color": "F4722B", "val": "single"}.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for name, params in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        border = OxmlElement(f"w:{name}")
        if params is None:
            border.set(qn("w:val"), "nil")
        else:
            border.set(qn("w:val"), params.get("val", "single"))
            border.set(qn("w:sz"), params.get("sz", size))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), params.get("color", BORDER_SOFT))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_run(paragraph, text, *, bold=False, size_pt=10, color_hex=None,
             font_name="Helvetica", caps=False, italic=False):
    from docx.shared import Pt
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_hex:
        run.font.color.rgb = _hex_to_rgb(color_hex)
    if caps:
        run.font.all_caps = True
    return run


def _new_para(container, *, alignment=None, space_before=0, space_after=2):
    from docx.shared import Pt
    if hasattr(container, "add_paragraph"):
        p = container.add_paragraph()
    else:
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ────────────────────────────────────────────────────────────────────
# Logo
# ────────────────────────────────────────────────────────────────────
def _add_logo_to_cell(cell, workspace, *, max_width_cm: float = 4.5):
    """Insère le logo workspace dans une cellule, ou laisse vide si absent."""
    from docx.shared import Cm
    if not workspace or not getattr(workspace, "logo", None):
        return False
    try:
        with workspace.logo.open("rb") as fh:
            data = fh.read()
        if not data:
            return False
        bio = io.BytesIO(data)
        para = cell.paragraphs[0]
        run = para.add_run()
        run.add_picture(bio, width=Cm(max_width_cm))
        return True
    except Exception as exc:
        logger.warning("Impossible d'insérer le logo workspace : %s", exc)
        return False


# ────────────────────────────────────────────────────────────────────
# Entrée principale
# ────────────────────────────────────────────────────────────────────
def render_invoice_docx(invoice, *, request=None) -> bytes:
    """Rend la facture en bytes .docx prêts à télécharger."""
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    workspace = invoice.workspace
    accent = (getattr(workspace, "accent_color", None) or f"#{ACCENT_DEFAULT}").lstrip("#")
    currency = invoice.currency or ""
    lines = list(invoice.lines.all().order_by("position", "id"))
    client = invoice.client
    project = invoice.project

    doc = Document()

    # ─── Marges page (compactes, comme le PDF) ───
    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    # ─── Style par défaut ───
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(9)
    style.font.color.rgb = _hex_to_rgb(TEXT_PRIMARY)

    # ============================================================
    # 1. EN-TÊTE : Logo (à gauche) + Tagline (à droite)
    # ============================================================
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.autofit = False
    header_tbl.columns[0].width = Cm(8.5)
    header_tbl.columns[1].width = Cm(8.5)
    cell_logo, cell_tag = header_tbl.rows[0].cells
    cell_logo.width = Cm(8.5)
    cell_tag.width = Cm(8.5)
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_tag.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _add_logo_to_cell(cell_logo, workspace, max_width_cm=4.5)

    tagline = (getattr(workspace, "tagline", "") or "").strip()
    if tagline:
        p_tag = cell_tag.paragraphs[0]
        p_tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p_tag, tagline, bold=True, size_pt=8, color_hex=accent, caps=True)

    # Filet orange sous l'en-tête
    for cell in (cell_logo, cell_tag):
        _set_cell_borders(
            cell,
            bottom={"sz": "6", "color": accent, "val": "single"},
        )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============================================================
    # 2. TITRE : "Facture N° XXX" + badge statut (1 ligne)
    # ============================================================
    title_tbl = doc.add_table(rows=1, cols=2)
    title_tbl.autofit = False
    title_tbl.columns[0].width = Cm(12)
    title_tbl.columns[1].width = Cm(5)
    c_title, c_badge = title_tbl.rows[0].cells
    c_title.width = Cm(12)
    c_badge.width = Cm(5)

    p_title = c_title.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(0)
    _add_run(p_title, "Facture", bold=True, size_pt=16, color_hex=TEXT_PRIMARY)
    invoice_no = invoice.number or "(brouillon)"
    _add_run(p_title, f"   N° {invoice_no}", bold=True, size_pt=11, color_hex=accent)

    # Badge statut
    status_key = (invoice.status or "").upper()
    badge = BADGE_COLORS.get(status_key)
    if badge:
        _set_cell_bg(c_badge, badge["bg"])
        p_badge = c_badge.paragraphs[0]
        p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_badge.paragraph_format.space_before = Pt(0)
        p_badge.paragraph_format.space_after = Pt(0)
        _add_run(p_badge, badge["label"], bold=True, size_pt=8,
                 color_hex=badge["fg"], caps=True)
    else:
        c_badge.paragraphs[0].add_run("")

    # Sous-titre projet + période sur 1 ligne (gain vertical)
    bits = []
    if getattr(invoice, "title", None) and project:
        bits.append(invoice.title)
    if invoice.period_start or invoice.period_end:
        bits.append(
            f"Période : {_fmt_date(invoice.period_start)} → {_fmt_date(invoice.period_end)}"
        )
    if bits:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(1)
        p_sub.paragraph_format.space_after = Pt(4)
        _add_run(p_sub, " · ".join(bits), size_pt=8, color_hex=TEXT_MUTED, italic=True)

    # ============================================================
    # 3. BLOC CLIENT (table 1×1 sur largeur réduite, fond gris clair)
    # ============================================================
    client_tbl = doc.add_table(rows=1, cols=1)
    client_tbl.autofit = False
    client_tbl.columns[0].width = Cm(9)
    cell_client = client_tbl.rows[0].cells[0]
    cell_client.width = Cm(9)
    _set_cell_bg(cell_client, BG_SOFT)
    # Bordures fines tout autour pour bien matérialiser le bloc
    _set_cell_borders(
        cell_client,
        top={"sz": "2", "color": BORDER_SOFT, "val": "single"},
        bottom={"sz": "2", "color": BORDER_SOFT, "val": "single"},
        left={"sz": "2", "color": BORDER_SOFT, "val": "single"},
        right={"sz": "2", "color": BORDER_SOFT, "val": "single"},
    )

    p_lbl = cell_client.paragraphs[0]
    p_lbl.paragraph_format.space_after = Pt(1)
    _add_run(p_lbl, "Adressée à", bold=True, size_pt=7, color_hex=TEXT_MUTED, caps=True)

    if client:
        # Nom légal en gras
        p_name = cell_client.add_paragraph()
        p_name.paragraph_format.space_after = Pt(0)
        _add_run(
            p_name,
            client.legal_name or client.name or "",
            bold=True, size_pt=9, color_hex=TEXT_PRIMARY,
        )

        # Lignes adresse / fiscalité
        details = []
        if getattr(client, "tax_id", ""):
            details.append(f"ID fiscal : {client.tax_id}")
        if getattr(client, "address_line1", ""):
            details.append(client.address_line1)
        if getattr(client, "address_line2", ""):
            details.append(client.address_line2)
        cp_city = " ".join(
            x for x in [getattr(client, "postal_code", ""), getattr(client, "city", "")]
            if x
        ).strip()
        if cp_city:
            details.append(cp_city)
        if getattr(client, "country", ""):
            details.append(client.country)
        if getattr(client, "contact_name", ""):
            details.append(f"À l'attention de : {client.contact_name}")
        for txt in details:
            p_d = cell_client.add_paragraph()
            p_d.paragraph_format.space_after = Pt(0)
            _add_run(p_d, txt, size_pt=8.5, color_hex=TEXT_PRIMARY)
    else:
        p_none = cell_client.add_paragraph()
        _add_run(p_none, "Aucun client renseigné", size_pt=9,
                 color_hex=TEXT_FAINT, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============================================================
    # 4. BANDEAU INFO : 4 colonnes (Date / Échéance / Projet / Devise)
    # ============================================================
    info_tbl = doc.add_table(rows=2, cols=4)
    info_tbl.autofit = False
    col_w = Cm(4.25)
    for col in info_tbl.columns:
        col.width = col_w
    for cell in info_tbl.rows[0].cells + info_tbl.rows[1].cells:
        cell.width = col_w

    labels = ["Date d'émission", "Échéance",
              ("Projet" if project else "Référence"), "Devise"]
    if project:
        ref = project.name
    elif getattr(invoice, "title", None):
        ref = invoice.title
    else:
        ref = "Facture libre"
    values = [
        _fmt_date(invoice.issue_date),
        _fmt_date(invoice.due_date),
        ref,
        currency or "—",
    ]

    for i, (lbl, val) in enumerate(zip(labels, values)):
        cell_lbl = info_tbl.rows[0].cells[i]
        cell_val = info_tbl.rows[1].cells[i]
        p_l = cell_lbl.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(0)
        _add_run(p_l, lbl, size_pt=7, color_hex=TEXT_FAINT, caps=True)
        p_v = cell_val.paragraphs[0]
        p_v.paragraph_format.space_after = Pt(0)
        _add_run(p_v, str(val), bold=True, size_pt=9, color_hex=TEXT_PRIMARY)
        # Bordures haut/bas sur les 4 cellules (effet bandeau)
        _set_cell_borders(
            cell_lbl,
            top={"sz": "4", "color": BORDER_SOFT, "val": "single"},
        )
        _set_cell_borders(
            cell_val,
            bottom={"sz": "4", "color": BORDER_SOFT, "val": "single"},
        )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============================================================
    # 5. TABLEAU DES LIGNES (Désignation / Qté / PU HT / Total HT)
    # ============================================================
    lines_tbl = doc.add_table(rows=1 + max(len(lines), 1), cols=4)
    lines_tbl.autofit = False
    # Largeurs colonnes
    w_desc = Cm(9.4)
    w_qty = Cm(1.9)
    w_pu = Cm(2.85)
    w_tot = Cm(2.85)
    for col, w in zip(lines_tbl.columns, [w_desc, w_qty, w_pu, w_tot]):
        col.width = w

    # En-tête (fond noir, texte blanc, caps)
    headers = ["Désignation", "Qté", "Prix unitaire HT", "Total HT"]
    head_row = lines_tbl.rows[0]
    for idx, head in enumerate(headers):
        cell = head_row.cells[idx]
        cell.width = [w_desc, w_qty, w_pu, w_tot][idx]
        _set_cell_bg(cell, TEXT_PRIMARY)
        p = cell.paragraphs[0]
        p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if idx == 0
                       else WD_ALIGN_PARAGRAPH.RIGHT)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, head, bold=True, size_pt=8, color_hex="FFFFFF", caps=True)

    # Lignes
    if lines:
        for ridx, line in enumerate(lines, start=1):
            row = lines_tbl.rows[ridx]
            cell_d = row.cells[0]
            cell_q = row.cells[1]
            cell_p = row.cells[2]
            cell_t = row.cells[3]
            cell_d.width = w_desc
            cell_q.width = w_qty
            cell_p.width = w_pu
            cell_t.width = w_tot

            # Désignation : label en gras + description en gris dessous
            p_lbl = cell_d.paragraphs[0]
            p_lbl.paragraph_format.space_after = Pt(0)
            _add_run(p_lbl, line.label or "—", bold=True, size_pt=9,
                     color_hex=TEXT_PRIMARY)
            desc = (line.description or "").strip()
            if desc:
                # On strip les balises HTML simples si présentes
                import re
                desc_clean = re.sub(r"<[^>]+>", "", desc)
                for chunk in desc_clean.split("\n"):
                    if chunk.strip():
                        p_d = cell_d.add_paragraph()
                        p_d.paragraph_format.space_after = Pt(0)
                        _add_run(p_d, chunk.strip(), size_pt=8,
                                 color_hex=TEXT_MUTED)

            # Qté
            p_q = cell_q.paragraphs[0]
            p_q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_q.paragraph_format.space_after = Pt(0)
            _add_run(p_q, _fmt_amount(line.quantity), size_pt=9,
                     color_hex=TEXT_PRIMARY)

            # PU
            p_p = cell_p.paragraphs[0]
            p_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_p.paragraph_format.space_after = Pt(0)
            _add_run(p_p, _fmt_amount(line.unit_price), size_pt=9,
                     color_hex=TEXT_PRIMARY)

            # Total
            p_t = cell_t.paragraphs[0]
            p_t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_t.paragraph_format.space_after = Pt(0)
            _add_run(p_t, _fmt_amount(line.total_amount), size_pt=9,
                     color_hex=TEXT_PRIMARY, bold=True)

            # Bordure inférieure fine
            for c in (cell_d, cell_q, cell_p, cell_t):
                _set_cell_borders(
                    c, bottom={"sz": "2", "color": BORDER_SOFT, "val": "single"},
                )
    else:
        # Rangée "Aucune ligne"
        empty_row = lines_tbl.rows[1]
        cell_empty = empty_row.cells[0]
        cell_empty.merge(empty_row.cells[1]).merge(empty_row.cells[2]).merge(
            empty_row.cells[3]
        )
        p_e = cell_empty.paragraphs[0]
        p_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_e, "Aucune ligne", size_pt=9,
                 color_hex=TEXT_FAINT, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============================================================
    # 6. TOTAUX (à droite : 2 colonnes Libellé / Valeur)
    # ============================================================
    totals_rows = [
        ("Sous-total HT", _fmt_amount(invoice.subtotal_ht)),
    ]
    if invoice.discount_amount:
        totals_rows.append(("Remise", f"- {_fmt_amount(invoice.discount_amount)}"))
    totals_rows.append(
        (f"TVA ({_fmt_amount(invoice.tax_rate)} %)",
         _fmt_amount(invoice.tax_amount))
    )
    totals_rows.append(
        ("TOTAL TTC", _fmt_amount(invoice.total_ttc, currency))
    )
    if invoice.paid_amount:
        totals_rows.append(("Encaissé", _fmt_amount(invoice.paid_amount)))
        try:
            remaining = invoice.remaining_due
        except Exception:
            remaining = (invoice.total_ttc or 0) - (invoice.paid_amount or 0)
        totals_rows.append(("Reste dû", _fmt_amount(remaining, currency)))

    # On utilise une table à 3 colonnes : spacer | label | value
    totals_tbl = doc.add_table(rows=len(totals_rows), cols=3)
    totals_tbl.autofit = False
    totals_tbl.columns[0].width = Cm(11)
    totals_tbl.columns[1].width = Cm(3)
    totals_tbl.columns[2].width = Cm(3)

    for r, (lbl, val) in enumerate(totals_rows):
        row = totals_tbl.rows[r]
        row.cells[0].width = Cm(11)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(3)

        # Spacer
        row.cells[0].paragraphs[0].add_run("")

        # Détection ligne TOTAL TTC
        is_total = lbl == "TOTAL TTC"

        p_l = row.cells[1].paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_after = Pt(0)
        _add_run(
            p_l, lbl,
            bold=is_total,
            size_pt=10.5 if is_total else 9,
            color_hex=accent if is_total else TEXT_PRIMARY,
            caps=is_total,
        )

        p_v = row.cells[2].paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_v.paragraph_format.space_after = Pt(0)
        _add_run(
            p_v, val,
            bold=is_total,
            size_pt=10.5 if is_total else 9,
            color_hex=accent if is_total else TEXT_PRIMARY,
        )

        if is_total:
            for cell in (row.cells[1], row.cells[2]):
                _set_cell_borders(
                    cell,
                    top={"sz": "8", "color": accent, "val": "single"},
                )

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============================================================
    # 7. NOTES (facultatif)
    # ============================================================
    notes_text = (getattr(invoice, "notes", "") or "").strip()
    if notes_text:
        notes_tbl = doc.add_table(rows=1, cols=1)
        notes_tbl.autofit = False
        notes_tbl.columns[0].width = Cm(17)
        cell_n = notes_tbl.rows[0].cells[0]
        cell_n.width = Cm(17)
        _set_cell_bg(cell_n, BG_SOFT)
        _set_cell_borders(
            cell_n,
            left={"sz": "12", "color": accent, "val": "single"},
        )

        p_h = cell_n.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(1)
        _add_run(p_h, "Notes", bold=True, size_pt=7,
                 color_hex=TEXT_MUTED, caps=True)

        import re
        notes_clean = re.sub(r"<[^>]+>", "", notes_text)
        for chunk in notes_clean.split("\n"):
            if chunk.strip():
                p_c = cell_n.add_paragraph()
                p_c.paragraph_format.space_after = Pt(0)
                _add_run(p_c, chunk.strip(), size_pt=8.5,
                         color_hex=TEXT_PRIMARY)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============================================================
    # 8. MENTIONS LÉGALES / COORDONNÉES BANCAIRES
    # ============================================================
    legal_bits = []
    bank_name = getattr(workspace, "bank_name", "")
    iban = getattr(workspace, "bank_iban", "")
    bic = getattr(workspace, "bank_bic", "")
    if bank_name or iban or bic:
        bank_line = "Coordonnées bancaires :"
        if bank_name:
            bank_line += f" {bank_name}"
        if iban:
            bank_line += f"  IBAN {iban}"
        if bic:
            bank_line += f"  BIC {bic}"
        legal_bits.append(bank_line)

    legal_text = (getattr(workspace, "legal_mentions", "") or "").strip()
    if legal_text:
        legal_bits.append(legal_text)

    if legal_bits:
        for txt in legal_bits:
            p_lg = doc.add_paragraph()
            p_lg.paragraph_format.space_after = Pt(0)
            _add_run(p_lg, txt, size_pt=7.5, color_hex=TEXT_MUTED)

    # ============================================================
    # 9. FOOTER : 3 colonnes (raison sociale / adresse / contact)
    # ============================================================
    section = doc.sections[0]
    footer = section.footer
    footer_tbl = footer.add_table(rows=1, cols=3, width=Cm(17.4))
    footer_tbl.autofit = False
    for col in footer_tbl.columns:
        col.width = Cm(5.8)

    legal_name = getattr(workspace, "legal_name", "") or getattr(workspace, "name", "")
    rccm = getattr(workspace, "legal_rccm", "")
    cc = getattr(workspace, "legal_cc", "")
    tax_id = getattr(workspace, "legal_tax_id", "")
    addr1 = getattr(workspace, "address_line1", "")
    addr2 = getattr(workspace, "address_line2", "")
    postal = getattr(workspace, "postal_code", "")
    city = getattr(workspace, "city", "")
    country = getattr(workspace, "country", "")
    phone = getattr(workspace, "phone", "")
    email = getattr(workspace, "email", "")
    website = getattr(workspace, "website", "")

    col_left_lines = [legal_name]
    if rccm: col_left_lines.append(f"RCCM : {rccm}")
    if cc: col_left_lines.append(f"CC : {cc}")
    if tax_id: col_left_lines.append(f"NIF : {tax_id}")

    col_mid_lines = []
    if addr1: col_mid_lines.append(addr1)
    if addr2: col_mid_lines.append(addr2)
    if postal or city:
        col_mid_lines.append(f"{postal} {city}".strip())
    if country: col_mid_lines.append(country)

    col_right_lines = []
    if phone: col_right_lines.append(f"Tél. : {phone}")
    if email: col_right_lines.append(email)
    if website: col_right_lines.append(website)

    footer_cells = footer_tbl.rows[0].cells
    for cell, lines_list, align in zip(
        footer_cells,
        [col_left_lines, col_mid_lines, col_right_lines],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT],
    ):
        cell.width = Cm(5.8)
        _set_cell_borders(
            cell,
            top={"sz": "4", "color": accent, "val": "single"},
        )
        if not lines_list:
            cell.paragraphs[0].add_run("")
            continue
        first = True
        for txt in lines_list:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _add_run(p, txt, size_pt=7, color_hex=TEXT_PRIMARY)

    # ─── Sortie binaire ───
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
