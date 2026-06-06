"""
DevFlow — Export d'un MeetingRecording en Word / PDF / Email.

Ce module produit des comptes-rendus à partir d'un ``MeetingRecording``
(résumé IA + extractions + transcription) avec le branding workspace.

Pendant le partage par email, on récupère les destinataires :
  * ``internal_participants`` de la réunion parente
  * Adresses email parsées depuis ``external_participants``

Toutes les opérations sont scopées au workspace de la réunion (la vue
appelante s'assure que l'utilisateur a bien accès). Aucune fuite
cross-tenant possible — on lit uniquement les données liées au
``meeting`` rattaché au ``recording``.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Iterable, Optional

from django.conf import settings
from django.core.mail import EmailMessage
from django.template.loader import render_to_string
from django.utils import timezone

from project import models as dm

logger = logging.getLogger(__name__)


# Regex simple email — suffisant pour extraire des emails depuis du texte libre
_EMAIL_RE = re.compile(r"[\w\.\-+]+@[\w\.\-]+\.\w+")


def _accent_rgb(workspace):
    """Convertit la couleur d'accent workspace en RGBColor python-docx."""
    from docx.shared import RGBColor
    accent = (getattr(workspace, "accent_color", None) or "#F4722B").lstrip("#")
    try:
        return RGBColor(
            int(accent[0:2], 16),
            int(accent[2:4], 16),
            int(accent[4:6], 16),
        )
    except Exception:
        return RGBColor(0xF4, 0x72, 0x2B)


def _extractions_by_kind(recording: dm.MeetingRecording, kind: str) -> list:
    """Liste des extractions d'un certain Kind, triées par confidence décroissante."""
    return list(
        recording.ai_extractions
        .filter(kind=kind)
        .order_by("-confidence", "id")
    )


def _add_workspace_branding(doc, workspace):
    """Ajoute en-tête (logo + tagline) et pied de page (coordonnées) au .docx."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    if workspace is None:
        return

    accent_rgb = _accent_rgb(workspace)
    section = doc.sections[0]

    # En-tête (logo + tagline)
    header = section.header
    ht = header.paragraphs[0]
    ht.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if getattr(workspace, "logo", None):
        try:
            run = ht.add_run()
            run.add_picture(workspace.logo.path, height=Cm(2.0))
        except Exception:
            pass
    if getattr(workspace, "tagline", None):
        ht_tag = header.add_paragraph(workspace.tagline)
        ht_tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in ht_tag.runs:
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = accent_rgb

    # Pied de page (raison sociale + coordonnées)
    footer = section.footer
    footer_lines = []
    if getattr(workspace, "legal_name", None) or getattr(workspace, "name", None):
        footer_lines.append(workspace.legal_name or workspace.name)
    if getattr(workspace, "legal_rccm", None):
        footer_lines.append(f"RCCM : {workspace.legal_rccm}")
    if getattr(workspace, "legal_cc", None):
        footer_lines.append(f"CC : {workspace.legal_cc}")
    if getattr(workspace, "address_line1", None):
        addr = workspace.address_line1
        if getattr(workspace, "address_line2", None):
            addr += " — " + workspace.address_line2
        footer_lines.append(addr)
    if getattr(workspace, "postal_code", None) or getattr(workspace, "city", None):
        footer_lines.append(
            f"{workspace.postal_code or ''} {workspace.city or ''}".strip()
        )
    if getattr(workspace, "phone", None):
        footer_lines.append(f"Tél. : {workspace.phone}")
    if getattr(workspace, "email", None):
        footer_lines.append(workspace.email)
    if not footer_lines:
        return
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in footer_lines:
        r = fp.add_run(line + "\n")
        r.font.size = Pt(8)


def render_recording_docx(recording: dm.MeetingRecording) -> bytes:
    """
    Produit un compte-rendu Word pour un ``MeetingRecording``.

    Structure :
      1. Branding workspace (logo + tagline + pied)
      2. Titre + bloc info réunion (titre, date, durée, lieu)
      3. Participants internes
      4. Résumé exécutif IA (summary_markdown rendu en paragraphes)
      5. Décisions détectées (extractions Kind.DECISION)
      6. Actions à mener (extractions Kind.ACTION) avec @owner + due
      7. Points de vigilance (extractions Kind.RISK)
      8. Transcription complète (final_transcript ou full_transcript)
      9. Pied : raison sociale + coordonnées

    Lève ImportError si ``python-docx`` n'est pas installé.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    meeting = recording.meeting
    workspace = recording.workspace
    accent_rgb = _accent_rgb(workspace)

    doc = Document()
    _add_workspace_branding(doc, workspace)

    # ── Titre ────────────────────────────────────────────────────────
    head = doc.add_paragraph()
    run = head.add_run("COMPTE-RENDU D'ENREGISTREMENT")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = accent_rgb

    h1 = doc.add_paragraph()
    h1.add_run(meeting.title if meeting else "Réunion sans titre").bold = True
    h1.runs[0].font.size = Pt(20)

    # ── Bloc info ────────────────────────────────────────────────────
    info = doc.add_paragraph()
    if meeting:
        info.add_run(f"Type : {meeting.get_meeting_type_display()}\n").bold = True
        info.add_run(
            f"Date : {timezone.localtime(meeting.scheduled_at).strftime('%A %d %B %Y à %H:%M')}\n"
        )
        info.add_run(f"Durée prévue : {meeting.duration_minutes} min\n")
        if meeting.location:
            info.add_run(f"Lieu : {meeting.location}\n")
        if meeting.meeting_link:
            info.add_run(f"Lien : {meeting.meeting_link}\n")
    info.add_run(f"Durée audio : {int(recording.duration_seconds or 0) // 60} min ").bold = True
    if recording.transcription_provider:
        info.add_run(
            f"· Transcription : {recording.transcription_provider}\n"
        )
    if recording.summary_provider:
        info.add_run(f"· Synthèse IA : {recording.summary_provider}\n")

    # ── Participants ─────────────────────────────────────────────────
    if meeting:
        internal = list(meeting.internal_participants.all())
        if internal:
            doc.add_heading("Participants", level=2)
            p = doc.add_paragraph()
            p.add_run("Internes : ").bold = True
            p.add_run(", ".join(
                u.get_full_name() or u.get_username() for u in internal
            ))
        if meeting.external_participants:
            p = doc.add_paragraph()
            p.add_run("Externes : ").bold = True
            p.add_run(meeting.external_participants)

    # ── Résumé IA ────────────────────────────────────────────────────
    if recording.summary_markdown:
        doc.add_heading("Résumé exécutif", level=2)
        for line in recording.summary_markdown.split("\n"):
            stripped = line.strip()
            if stripped:
                # On retire la syntaxe markdown la plus basique (# ## - *)
                clean = re.sub(r"^[#\*\-\+\s]+", "", stripped)
                if clean:
                    doc.add_paragraph(clean)

    # ── Décisions détectées ──────────────────────────────────────────
    decisions = _extractions_by_kind(recording, dm.RecordingAIExtraction.Kind.DECISION)
    if decisions:
        doc.add_heading("Décisions", level=2)
        for d in decisions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(d.title).bold = True
            if d.description:
                p.add_run(" — " + d.description)
            if d.assignee_hint:
                p.add_run(f"  (resp. : {d.assignee_hint})")

    # ── Actions à mener ──────────────────────────────────────────────
    actions = _extractions_by_kind(recording, dm.RecordingAIExtraction.Kind.ACTION)
    if actions:
        doc.add_heading("Actions à mener", level=2)
        for a in actions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(a.title).bold = True
            details = []
            if a.assignee_hint:
                details.append(f"@{a.assignee_hint}")
            if a.due_date_hint:
                details.append(f"⌛ {a.due_date_hint}")
            if a.priority_hint:
                details.append(f"[{a.priority_hint}]")
            if details:
                p.add_run("  " + " · ".join(details))
            if a.description:
                p.add_run("\n    " + a.description).italic = True

    # ── Risques / vigilance ──────────────────────────────────────────
    risks = _extractions_by_kind(recording, dm.RecordingAIExtraction.Kind.RISK)
    if risks:
        doc.add_heading("Points de vigilance", level=2)
        for r in risks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(r.title).bold = True
            if r.description:
                p.add_run(" — " + r.description)

    # ── Transcription verbatim (annexe) ──────────────────────────────
    transcript = recording.final_transcript or recording.full_transcript
    if transcript:
        doc.add_paragraph()
        doc.add_heading("Transcription complète", level=2)
        # On découpe en paragraphes par double saut de ligne
        for block in transcript.split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_recording_pdf(recording: dm.MeetingRecording) -> bytes:
    """
    Produit un PDF du compte-rendu via WeasyPrint, avec branding workspace.

    Réutilise le résolveur de logo en base64 (``invoice_pdf._resolve_logo_uri``)
    pour éviter les fetches réseau de WeasyPrint.
    """
    from weasyprint import HTML
    from project.services.invoice_pdf import _resolve_logo_uri

    meeting = recording.meeting
    workspace = recording.workspace
    logo_uri = _resolve_logo_uri(workspace, request=None) if workspace else ""

    ctx = {
        "recording": recording,
        "meeting": meeting,
        "workspace": workspace,
        "logo_uri": logo_uri,
        "internal_participants": (
            list(meeting.internal_participants.all()) if meeting else []
        ),
        "decisions": _extractions_by_kind(
            recording, dm.RecordingAIExtraction.Kind.DECISION,
        ),
        "actions": _extractions_by_kind(
            recording, dm.RecordingAIExtraction.Kind.ACTION,
        ),
        "risks": _extractions_by_kind(
            recording, dm.RecordingAIExtraction.Kind.RISK,
        ),
        "transcript": recording.final_transcript or recording.full_transcript or "",
    }
    html_str = render_to_string("project/meeting/recording_pdf.html", ctx)
    out = io.BytesIO()
    HTML(string=html_str).write_pdf(target=out)
    return out.getvalue()


def _collect_recipients(
    recording: dm.MeetingRecording,
    *,
    include_external: bool = True,
    extra_emails: Optional[Iterable[str]] = None,
) -> list[str]:
    """Liste dédoublonnée des destinataires email pour un recording."""
    recipients: list[str] = []
    meeting = recording.meeting
    if meeting:
        for u in meeting.internal_participants.all():
            if u.email and u.email not in recipients:
                recipients.append(u.email)
        if include_external and meeting.external_participants:
            for line in meeting.external_participants.splitlines():
                for match in _EMAIL_RE.findall(line):
                    if match not in recipients:
                        recipients.append(match)
    for e in extra_emails or []:
        e = (e or "").strip()
        if e and e not in recipients:
            recipients.append(e)
    return recipients


def send_recording_email(
    recording: dm.MeetingRecording,
    *,
    include_external: bool = True,
    extra_emails: Optional[Iterable[str]] = None,
) -> int:
    """
    Envoie le .docx + un corps HTML (résumé IA) aux participants.

    Retourne le nombre d'emails effectivement envoyés (échecs muets).

    Sécurité : seules les adresses des participants de la réunion +
    `extra_emails` sont utilisées. Le sujet et le corps respectent le
    branding workspace dans la mesure du possible.
    """
    meeting = recording.meeting
    recipients = _collect_recipients(
        recording,
        include_external=include_external,
        extra_emails=extra_emails,
    )
    if not recipients:
        logger.info("Recording %s: no recipients for email", recording.pk)
        return 0

    # Génère le .docx
    docx_bytes = None
    try:
        docx_bytes = render_recording_docx(recording)
    except Exception as exc:
        logger.warning(
            "Cannot render docx for recording %s: %s", recording.pk, exc
        )

    # Corps texte plain + HTML
    body_text = recording.summary_markdown or (
        "Veuillez trouver en pièce jointe le compte-rendu de la réunion "
        f"« {meeting.title if meeting else '—'} ».\n\n"
        "— Compte-rendu DevFlow"
    )

    meeting_date = (
        timezone.localtime(meeting.scheduled_at).strftime("%d/%m/%Y")
        if meeting else timezone.localtime(recording.created_at).strftime("%d/%m/%Y")
    )
    subject = f"[CR] {meeting.title if meeting else 'Réunion'} — {meeting_date}"

    from_email = (
        getattr(settings, "DEFAULT_FROM_EMAIL", None)
        or getattr(settings, "EMAIL_HOST_USER", None)
        or "noreply@devflow.local"
    )

    sent = 0
    for to in recipients:
        try:
            msg = EmailMessage(
                subject=subject,
                body=body_text,
                from_email=from_email,
                to=[to],
            )
            if docx_bytes:
                # Nom de fichier propre
                safe_title = re.sub(
                    r"[^A-Za-z0-9\-_]+", "_",
                    meeting.title if meeting else "Recording",
                )
                filename = f"CR-{safe_title}-{meeting_date.replace('/', '-')}.docx"
                msg.attach(
                    filename, docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            msg.send(fail_silently=True)
            sent += 1
        except Exception as exc:
            logger.warning("Recording email send failed to %s: %s", to, exc)
    return sent
